import os
import moviepy.editor
import speech_recognition as sr

video=moviepy.editor.VideoFileClip('/content/Final_summ/code/1_1.mp4')

audio=video.audio

audio.write_audiofile('/content/Final_summ/code/1.wav')

filename="1.wav"

r=sr.Recognizer()

f=open("myfile.txt","w")

with sr.AudioFile(filename) as source:
	audio_data=r.record(source)
	text=r.recognize_google(audio_data)
	f.write(text)
f.close()

from docx import Document
#from docx.text.parargaph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
#import glob
 
doc=Document()
styles = doc.styles
style = styles.add_style('Tahoma', WD_STYLE_TYPE.PARAGRAPH) #Tahoma is the name I set because that's the font I'm gonna use
style.font.size = Pt(11)

file="myfile.txt"
        
with open(file, 'r', encoding='UTF-8') as openfile:
    line=openfile.read()
    t=doc.add_heading("YouTube Video Summary")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    p=doc.add_paragraph("\t"+line.capitalize()+".")
    p.style = doc.styles['Tahoma']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(file + ".docx")
    
os.system(file + ".docx")
os.remove(1.wav)
os.remove(myfile.txt)
os.remove(data)
